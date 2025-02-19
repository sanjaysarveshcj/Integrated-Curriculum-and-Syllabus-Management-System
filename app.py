from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google_auth_oauthlib.flow import Flow
import os
import json
import tempfile
from docx import Document
import docx.enum.text
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
import io
import re
import logging
from datetime import timedelta, datetime
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import tempfile
from dotenv import load_dotenv
from google_auth_oauthlib.flow import Flow
from models import db, User, DriveDirectory, DocumentApproval
from drive_helper import create_directory_structure, get_viewable_folder_id, get_google_drive_service
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.permanent_session_lifetime = timedelta(minutes=30)

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Initialize SQLAlchemy with app
db.init_app(app)

# Load environment variables
load_dotenv()

# Allow OAuth2 insecure transport for development
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-here')
app.config['SESSION_TYPE'] = 'filesystem'

CURRICULUM_FOLDER_ID = "11efeP3LJ23w2lFt1AJI_jJBNyseRHPfn"  # Your main folder ID

# Document manipulation functions
def clean_text(value):
    """Ensures the input is a string before processing."""
    if value is None:
        return ""
    return str(value).strip()

def clean_pdf_text(text):
    """Cleans text that might have been copied from a PDF."""
    if not text:
        return ""
    # Remove extra whitespace and normalize line endings
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove any special characters or control characters
    text = ''.join(char for char in text if char.isprintable())
    return text

def clean_int(value):
    """Converts a string to integer, returns 0 if invalid."""
    try:
        return int(value)
    except (ValueError, TypeError):
        return 0

def replace_list_section(doc, placeholder, items, title=""):
    """
    Replaces a placeholder with a properly formatted numbered list while keeping the content at the correct position.
    - `placeholder`: The placeholder text to replace (e.g., `{Objectives}`)
    - `items`: The list of items to insert
    - `title`: The title of the section (optional)
    """
    for i, paragraph in enumerate(doc.paragraphs):
        if placeholder in paragraph.text:
            parent = paragraph._element.getparent()  # Get parent XML element
            paragraph.text = ""  # Clear placeholder but keep paragraph position
            if not items:
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                return 
            # Preserve the document's original paragraph format
            paragraph_format = paragraph.paragraph_format
            
            # Insert title (if provided)
            if title:
                title_paragraph = paragraph.insert_paragraph_before()
                title_paragraph.style = paragraph.style  # Keep same style
                title_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  # Maintain document indentation
                title_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  # Keep first-line formatting
                
                title_run = title_paragraph.add_run(title)
                title_run.bold = True  
                title_run.font.size = Pt(11)

            # Insert list items directly after the placeholder
            for index, item in enumerate(items, 1):
                item_paragraph = paragraph.insert_paragraph_before("")
                item_paragraph.style = paragraph.style  # Keep same style
                item_paragraph.paragraph_format.left_indent = paragraph_format.left_indent  # Maintain document indentation
                item_paragraph.paragraph_format.first_line_indent = paragraph_format.first_line_indent  # Keep first-line formatting
                
                # Manually add numbering (bold)
                item_run = item_paragraph.add_run(f"{index}. ")
                item_run.bold = True  
                item_run.font.size = Pt(11)
                
                # Add the actual content
                content_run = item_paragraph.add_run(item.strip())  
                content_run.bold = False  
                content_run.font.size = Pt(11)

                # Set paragraph indentation
                pPr = item_paragraph._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "645")  # Use document's original left indentation
                ind.set(qn("w:hanging"), "365")  # Hanging indent for text (0.25 inch)
                pPr.append(ind)

            return  # Stop after first occurrence

def replace_course_name_in_table(doc, course_name):
    """Replaces {CourseName} in the document's tables while maintaining formatting."""
    placeholder = "{CourseName}"
    value = course_name if course_name else "<REMOVE>"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

def replace_course_code_in_table(doc, course_code):
    """Replaces {CourseCode} in the document's tables while maintaining formatting."""
    placeholder = "{CourseCode}"
    value = course_code if course_code else "<REMOVE>"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

def replace_semester(doc, semester):
    """Replaces {Semester} while maintaining formatting."""
    placeholder = "{Semester}"
    value = semester if semester else "<REMOVE>"
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)

def replace_course_description(doc, course_description):
    """Replaces {CourseDescription} while maintaining formatting and indentation."""
    placeholder = "{CourseDescription}"
    title = "COURSE DESCRIPTION" if course_description else "<REMOVE>"
    value = course_description if course_description else "<REMOVE>"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            if value == "<REMOVE>":
                # Remove the paragraph if there's no content
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            else:
                # Add title in bold
                title_paragraph = paragraph.insert_paragraph_before()
                title_run = title_paragraph.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(11)
                
                # Add content
                paragraph.text = value
                for run in paragraph.runs:
                    run.bold = False
                    run.font.size = Pt(11)

def replace_prerequisites(doc, prerequisites):
    """Replaces {Prerequisites} while maintaining formatting."""
    placeholder = "{Prerequisites}"
    title = "PREREQUISITES" if prerequisites else "<REMOVE>"
    value = prerequisites if prerequisites else "<REMOVE>"

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            if value == "<REMOVE>":
                # Remove the paragraph if there's no content
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            else:
                # Add title in bold
                title_paragraph = paragraph.insert_paragraph_before()
                title_run = title_paragraph.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(11)
                
                # Add content
                paragraph.text = value
                for run in paragraph.runs:
                    run.bold = False
                    run.font.size = Pt(11)

def replace_course_format(doc, course_format):
    """Replaces {CourseFormat} while maintaining formatting."""
    placeholder = "{CourseFormat}"
    value = course_format if course_format else "<REMOVE>"
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            if value == "<REMOVE>":
                # Remove the paragraph if there's no content
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            else:
                paragraph.text = value
                for run in paragraph.runs:
                    run.bold = False
                    run.font.size = Pt(11)

def replace_assessments_grading(doc, assessments_grading):
    """Replaces {AssessmentsGrading} while maintaining formatting."""
    placeholder = "{AssessmentsGrading}"
    value = assessments_grading if assessments_grading else "<REMOVE>"
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            if value == "<REMOVE>":
                # Remove the paragraph if there's no content
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            else:
                paragraph.text = value
                for run in paragraph.runs:
                    run.bold = False
                    run.font.size = Pt(11)

def replace_practical_periods(doc, practical_periods):
    """Replaces {PracticalPeriods} while maintaining formatting."""
    placeholder = "{PracticalPeriods}"
    value = practical_periods if practical_periods else "<REMOVE>"
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            if value == "<REMOVE>":
                # Remove the paragraph if there's no content
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
            else:
                paragraph.text = value
                for run in paragraph.runs:
                    run.bold = False
                    run.font.size = Pt(11)

def format_course_outcomes(doc, placeholder, outcomes):
    """Formats course outcomes with proper numbering and indentation."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            if not outcomes:
                # Remove the paragraph if there are no outcomes
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                return

            # Add title
            title_paragraph = paragraph.insert_paragraph_before()
            title_run = title_paragraph.add_run("COURSE OUTCOMES")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # Add outcomes with proper formatting
            for index, outcome in enumerate(outcomes, 1):
                outcome_paragraph = paragraph.insert_paragraph_before("")
                
                # Add the outcome number (CO1, CO2, etc.)
                outcome_run = outcome_paragraph.add_run(f"CO{index}: ")
                outcome_run.bold = True
                outcome_run.font.size = Pt(11)
                
                # Add the outcome text
                text_run = outcome_paragraph.add_run(outcome.strip())
                text_run.bold = False
                text_run.font.size = Pt(11)

                # Set paragraph indentation
                pPr = outcome_paragraph._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "645")
                ind.set(qn("w:hanging"), "365")
                pPr.append(ind)

            # Remove the original placeholder paragraph
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
            return

def replace_units_with_formatting(doc, units):
    """Formats units with proper numbering, periods, and indentation."""
    for paragraph in doc.paragraphs:
        if "{Units}" in paragraph.text:
            if not units:
                # Remove the paragraph if there are no units
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                return

            # Add title
            title_paragraph = paragraph.insert_paragraph_before()
            title_run = title_paragraph.add_run("COURSE CONTENT")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # Add units with proper formatting
            for index, unit in enumerate(units, 1):
                # Unit title with periods
                unit_title_paragraph = paragraph.insert_paragraph_before()
                title_run = unit_title_paragraph.add_run(f"UNIT {index}: {unit['title']} ({unit['periods']} Periods)")
                title_run.bold = True
                title_run.font.size = Pt(11)

                # Unit content
                content_paragraph = paragraph.insert_paragraph_before()
                content_run = content_paragraph.add_run(unit['content'])
                content_run.bold = False
                content_run.font.size = Pt(11)

                # Set paragraph indentation for content
                pPr = content_paragraph._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "645")
                pPr.append(ind)

            # Remove the original placeholder paragraph
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
            return

def replace_total_periods(doc, units):
    """Updates the total periods in the document."""
    total_periods = sum(unit['periods'] for unit in units)
    placeholder = "{TotalPeriods}"
    
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(total_periods))

def replace_youtube_references_with_formatting(doc, youtube_references):
    """Formats YouTube references with proper numbering and indentation."""
    for paragraph in doc.paragraphs:
        if "{YouTubeReferences}" in paragraph.text:
            if not youtube_references:
                # Remove the paragraph if there are no references
                p_element = paragraph._element
                p_element.getparent().remove(p_element)
                return

            # Add title
            title_paragraph = paragraph.insert_paragraph_before()
            title_run = title_paragraph.add_run("YOUTUBE REFERENCES")
            title_run.bold = True
            title_run.font.size = Pt(11)

            # Add references with proper formatting
            for index, (title, desc, url) in enumerate(youtube_references, 1):
                # Reference number and title
                ref_paragraph = paragraph.insert_paragraph_before("")
                ref_run = ref_paragraph.add_run(f"{index}. {title}")
                ref_run.bold = True
                ref_run.font.size = Pt(11)

                # Description
                desc_paragraph = paragraph.insert_paragraph_before("")
                desc_run = desc_paragraph.add_run(desc)
                desc_run.bold = False
                desc_run.font.size = Pt(11)

                # URL
                url_paragraph = paragraph.insert_paragraph_before("")
                url_run = url_paragraph.add_run(url)
                url_run.bold = False
                url_run.font.size = Pt(11)
                url_run.underline = True

                # Set paragraph indentation
                for p in [ref_paragraph, desc_paragraph, url_paragraph]:
                    pPr = p._element.get_or_add_pPr()
                    ind = OxmlElement("w:ind")
                    ind.set(qn("w:left"), "645")
                    pPr.append(ind)

            # Remove the original placeholder paragraph
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.route('/')
def index():
    return render_template('login.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        user = User.query.filter_by(email=email).first()
        
        if user and check_password_hash(user.password_hash, password):
            login_user(user, remember=True)
            session.permanent = True  # Make the session permanent
            
            if user.role == 'principal':
                return redirect(url_for('principal_dashboard'))
            elif user.role == 'hod':
                return redirect(url_for('hod_dashboard'))
            elif user.role == 'teacher':
                return redirect(url_for('teacher_dashboard'))
            elif user.role == 'advisor':
                return redirect(url_for('advisor_dashboard'))
        
        flash('Invalid email or password')
    return render_template('login.html')

@app.route('/principal_dashboard')
@login_required
def principal_dashboard():
    if current_user.role != 'principal':
        return redirect(url_for('login'))
    hods = User.query.filter_by(role='hod').all()
    return render_template('principal_dashboard.html', hods=hods)

@app.route('/hod_dashboard')
@login_required
def hod_dashboard():
    if current_user.role != 'hod':
        flash('Access denied')
        return redirect(url_for('index'))
        
    # Get staff members under this HOD
    staff = User.query.filter_by(
        department=current_user.department,
        created_by=current_user.id
    ).all()
    
    # Get department folder ID
    dept_folder = DriveDirectory.query.filter_by(
        department=current_user.department,
        type='department'
    ).first()
    
    folder_id = dept_folder.drive_id if dept_folder else None
    
    return render_template(
        'hod_dashboard.html',
        staff=staff,
        folder_id=folder_id
    )

@app.route('/advisor_dashboard')
@login_required
def advisor_dashboard():
    if current_user.role != 'advisor':
        flash('Access denied. You must be an advisor.', 'error')
        return redirect(url_for('index'))

    # Get regular document approvals
    approval_docs = DocumentApproval.query.filter_by(
        department=current_user.department,
        semester=current_user.semester,
        document_type='regular'
    ).all()

    # Get HOD approval documents
    hod_approval_docs = DocumentApproval.query.filter_by(
        department=current_user.department,
        semester=current_user.semester,
        document_type='hod'
    ).all()

    # Get syllabus approval documents
    syllabus_approval_docs = DocumentApproval.query.filter_by(
        department=current_user.department,
        semester=current_user.semester,
        document_type='syllabus'
    ).all()

    return render_template('advisor_dashboard.html', 
                         approval_docs=approval_docs,
                         hod_approval_docs=hod_approval_docs,
                         syllabus_approval_docs=syllabus_approval_docs)

@app.route('/teacher_dashboard')
@login_required
def teacher_dashboard():
    if current_user.role != 'teacher':
        flash('Access denied')
        return redirect(url_for('index'))
    
    # Get the subjects folder ID for the current user
    subjects_folder = DriveDirectory.query.filter_by(
        department=current_user.department,
        type='subject',
        semester=current_user.semester
    ).first()
    
    folder_id = subjects_folder.drive_id if subjects_folder else None
    
    # Get approval documents
    approval_docs = DocumentApproval.query.filter_by(
        department=current_user.department,
        semester=str(current_user.semester)
    ).all()
    
    return render_template('teacher_dashboard.html', folder_id=folder_id, approval_docs=approval_docs)

@app.route('/front_form')
@login_required
def front_form_redirect():
    return redirect(url_for('frontform'))

@app.route('/frontform', methods=['GET', 'POST']) 
@login_required
def frontform():
    if request.method == "POST":
        try:
            # Get the subjects folder ID for the current user
            subjects_folder = DriveDirectory.query.filter_by(
                department=current_user.department,
                type='subject',
                semester=current_user.semester
            ).first()
            
            if not subjects_folder:
                flash('Subject folder not found. Please contact your HOD.')
                return redirect(url_for('advisor_dashboard'))
            
            form_data = process_form_data(request.form)
            doc = generate_doc(form_data, subjects_folder.drive_id)
            return doc

        except Exception as e:
            flash(f'Error generating document: {str(e)}')
            return redirect(url_for('frontform'))
    return render_template('frontform.html')

@app.route('/syllabusform', methods=['GET', 'POST'])
@login_required
def syllabusform():
    if request.method == 'POST':
        temp_file = None
        try:
            # Get the subjects folder ID for the current user
            subjects_folder = DriveDirectory.query.filter_by(
                department=current_user.department,
                type='subject',
                semester=current_user.semester
            ).first()
            
            if not subjects_folder:
                flash('Subject folder not found. Please contact your HOD.')
                if current_user.role == 'teacher':
                    return redirect(url_for('teacher_dashboard'))
                else:
                    return redirect(url_for('advisor_dashboard'))

            template_path = os.path.join(BASE_DIR, "template.docx")
            if not os.path.exists(template_path):
                return "Error: template.docx not found!", 404

            doc = Document(template_path)

            # Collect form data and clean all text fields if pasted from PDF
            semester = request.form.get('Semester', '')
            course_name = request.form.get('CourseName', '')
            course_code = request.form.get('CourseCode', '')
            course_description = clean_pdf_text(request.form.get('CourseDescription', ''))
            prerequisites = clean_pdf_text(request.form.get('Prerequisites', ''))
            objectives = [clean_pdf_text(obj.strip()) for obj in request.form.getlist('objective') if obj.strip()]
            experiments = [clean_pdf_text(obj.strip()) for obj in request.form.getlist('experiments') if obj.strip()]
            course_outcomes = [clean_pdf_text(outcome.strip()) for outcome in request.form.getlist('course_outcome') if outcome.strip()]
            textbooks = [clean_pdf_text(textbook.strip()) for textbook in request.form.getlist('textbook') if textbook.strip()]
            references = [clean_pdf_text(reference.strip()) for reference in request.form.getlist('reference') if reference.strip()]
            assessments_grading = clean_pdf_text(request.form.get('AssessmentsGrading', ''))
            course_format = clean_pdf_text(request.form.get('courseformat', ''))
            assessments = clean_pdf_text(request.form.get('assessments', ''))
            grading = clean_pdf_text(request.form.get('grading', ''))
            practical_periods = request.form.get('practical_periods', '')

            # Process units data
            units = []
            i = 1
            while True:
                title = request.form.get(f'unit_title_{i}', '')
                content = request.form.get(f'unit_content_{i}', '')
                periods = request.form.get(f'unit_periods_{i}', '')

                if not title or not content or not periods:
                    break

                units.append({
                    'title': clean_pdf_text(title),
                    'content': clean_pdf_text(content),
                    'periods': clean_int(periods)
                })
                i += 1

            # Process YouTube references
            youtube_references = []
            i = 1
            while True:
                youtube_title = request.form.get(f'youtube_title_{i}', '')
                youtube_desc = request.form.get(f'youtube_desc_{i}', '')
                youtube_url = request.form.get(f'youtube_url_{i}', '')

                if not youtube_title or not youtube_desc or not youtube_url:
                    break

                youtube_references.append((youtube_title, youtube_desc, youtube_url))
                i += 1

            # Apply all replacements
            replace_list_section(doc, "{Objectives}", objectives, title="COURSE OBJECTIVES")
            replace_list_section(doc, "{Experiments}", experiments, title="LIST OF EXPERIMENTS")
            replace_list_section(doc, "{Textbooks}", textbooks, title="TEXTBOOKS")
            replace_list_section(doc, "{References}", references, title="REFERENCES")
            format_course_outcomes(doc, "{CourseOutcomes}", course_outcomes)
            replace_units_with_formatting(doc, units)
            replace_semester(doc, semester)
            replace_course_name_in_table(doc, course_name)
            replace_course_code_in_table(doc, course_code)
            replace_course_description(doc, course_description)
            replace_prerequisites(doc, prerequisites)
            replace_course_format(doc, course_format)
            replace_assessments_grading(doc, assessments_grading)
            replace_practical_periods(doc, practical_periods)
            replace_youtube_references_with_formatting(doc, youtube_references)
            replace_total_periods(doc, units)

            # Create a temporary file with a unique name
            temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_file_path = temp_file.name
            temp_file.close()  # Close the file handle immediately

            # Save the document to the temporary file
            doc.save(temp_file_path)
                
            # Get Drive API service
            drive_service, auth_url = get_google_drive_service()
            if not drive_service:
                if auth_url:
                    session['next_url'] = request.url  # Save current URL to redirect back after auth
                    return redirect(auth_url)
                flash('Failed to get Google Drive service. Please try again.')
                return redirect(request.referrer)
            
            try:
                # Upload DOCX file and convert to Google Docs format
                file_metadata = {
                    'name': f"{request.form.get('CourseName', 'document')}_{request.form.get('Semester', '')}_Syllabus",
                    'parents': [subjects_folder.drive_id],
                    'mimeType': 'application/vnd.google-apps.document'
                }
                
                media = MediaFileUpload(
                    temp_file_path,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    resumable=True
                )
                
                file = drive_service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields='id'
                ).execute()
                
                # Redirect to the Google Docs URL
                google_docs_url = f'https://docs.google.com/document/d/{file.get("id")}/edit'
                return redirect(google_docs_url)

            finally:
                # Clean up: Remove the temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception as e:
                    logging.error(f"Error removing temporary file: {e}")

        except Exception as e:
            flash(f'Error generating document: {str(e)}')
            return redirect(request.referrer)
        finally:
            # Additional cleanup in case the file wasn't removed
            if temp_file and os.path.exists(temp_file.name):
                try:
                    os.unlink(temp_file.name)
                except Exception as e:
                    logging.error(f"Error removing temporary file in finally block: {e}")

    return render_template('syllabusform.html')

@app.route('/create_user', methods=['POST'])
@login_required
def create_user():
    if current_user.role == 'principal' and request.form.get('role') == 'hod':
        username = request.form.get('username')
        password = request.form.get('password')
        email = request.form.get('email')
        department = request.form.get('department')
        
        if User.query.filter_by(username=username).first():
            flash('Username already exists')
            return redirect(url_for('principal_dashboard'))
            
        if User.query.filter_by(email=email).first():
            flash('Email already exists')
            return redirect(url_for('principal_dashboard'))
            
        hod = User(
            username=username,
            email=email,
            role='hod',
            department=department,
            created_by=current_user.id
        )
        hod.set_password(password)
        db.session.add(hod)
        db.session.commit()
        flash('HOD account created successfully')
        return redirect(url_for('principal_dashboard'))
        
    elif current_user.role == 'hod':
        username = request.form.get('username')
        password = request.form.get('password')
        email = request.form.get('email')
        role = request.form.get('role')
        semester = request.form.get('semester')
        
        if User.query.filter_by(username=username).first():
            flash('Username already exists')
            return redirect(url_for('hod_dashboard'))
            
        if User.query.filter_by(email=email).first():
            flash('Email already exists')
            return redirect(url_for('hod_dashboard'))
            
        staff = User(
            username=username,
            email=email,
            role=role,
            department=current_user.department,
            semester=semester,
            created_by=current_user.id
        )
        staff.set_password(password)
        db.session.add(staff)
        db.session.commit()
        flash('Staff account created successfully')
        return redirect(url_for('hod_dashboard'))
        
    return redirect(url_for('principal_dashboard'))

@app.route('/delete_user/<int:user_id>')
@login_required
def delete_user(user_id):
    user_to_delete = User.query.get_or_404(user_id)
    
    # Only principal can delete HODs
    if current_user.role == 'principal' and user_to_delete.role == 'hod':
        db.session.delete(user_to_delete)
        db.session.commit()
        flash('HOD account deleted successfully')
        return redirect(url_for('principal_dashboard'))
    
    # Only HODs can delete their created staff
    elif current_user.role == 'hod' and user_to_delete.created_by == current_user.id:
        db.session.delete(user_to_delete)
        db.session.commit()
        flash('Staff account deleted successfully')
        return redirect(url_for('hod_dashboard'))
    
    flash('You do not have permission to delete this user')
    return redirect(request.referrer)

@app.route('/principal/create_directory', methods=['GET', 'POST'])
@login_required
def create_directory():
    if current_user.role != 'principal':
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        department = request.form.get('department')
        regulation_code = request.form.get('regulation_code')
        
        try:
            service, auth_url = get_google_drive_service()
            if not service:
                # Store the form data in session
                session['pending_directory'] = {
                    'department': department,
                    'regulation_code': regulation_code
                }
                return redirect(auth_url)
            
            create_directory_structure(service, department, regulation_code, CURRICULUM_FOLDER_ID)
            flash('Directory structure created successfully!', 'success')
        except Exception as e:
            flash(f'Error: {str(e)}', 'error')
            
        return redirect(url_for('create_directory'))
        
    return render_template('create_directory.html')

@app.route('/oauth2callback')
def oauth2callback():
    try:
        # Get the Flow instance from the state
        flow = Flow.from_client_secrets_file(
            'credentials.json',
            scopes=['https://www.googleapis.com/auth/drive.file'],
            redirect_uri=url_for('oauth2callback', _external=True)
        )
        
        # Fetch the authorization code from the request
        authorization_response = request.url
        flow.fetch_token(authorization_response=authorization_response)
        
        # Save the credentials
        credentials = flow.credentials
        with open('token.json', 'w') as token:
            token.write(credentials.to_json())
        
        # Check if there was a pending upload
        if 'pending_upload' in session:
            logging.info('Found pending upload in session')
            return redirect(url_for('teacher_dashboard'))
        
        # Redirect back to the original page
        next_url = session.get('next_url')
        if next_url:
            session.pop('next_url')
            return redirect(next_url)
        return redirect(url_for('index'))
        
    except Exception as e:
        logging.error('Error during authentication: %s', str(e))
        flash(f'Error during authentication: {str(e)}')
        return redirect(url_for('index'))

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

def generate_doc(context, folder_id):
    """Generate a document and save it to Google Drive"""
    try:
        # Get the template based on document type
        template_path = os.path.join(os.path.dirname(__file__), 'templates', 'syllabus_template.docx')
        doc = DocxTemplate(template_path)
        
        # Render the template with the context
        doc.render(context)
        
        # Save the DOCX file to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            doc.save(temp_docx.name)
            
            # Get Drive API service
            creds = None
            if 'credentials' in session:
                creds = google.oauth2.credentials.Credentials(**session['credentials'])
            
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    flow = google_auth_oauthlib.flow.Flow.from_client_secrets_file(
                        'client_secrets.json',
                        scopes=['https://www.googleapis.com/auth/drive.file']
                    )
                    flow.redirect_uri = url_for('oauth2callback', _external=True)
                    
                    authorization_url, state = flow.authorization_url(
                        access_type='offline',
                        include_granted_scopes='true'
                    )
                    session['state'] = state
                    return redirect(authorization_url)
                    
                session['credentials'] = {
                    'token': creds.token,
                    'refresh_token': creds.refresh_token,
                    'token_uri': creds.token_uri,
                    'client_id': creds.client_id,
                    'client_secret': creds.client_secret,
                    'scopes': creds.scopes
                }
            
            # Create Drive API service
            drive_service = build('drive', 'v3', credentials=creds)
            
            # Upload DOCX file and convert to Google Docs format
            file_metadata = {
                'name': f"{context.get('course_name', 'document')}_{context.get('semester', '')}",
                'parents': [folder_id],
                'mimeType': 'application/vnd.google-apps.document'  # This converts to Google Docs format
            }
            
            media = MediaFileUpload(
                temp_docx.name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                resumable=True
            )
            
            file = drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id'
            ).execute()
            
            # Clean up the temporary file
            os.unlink(temp_docx.name)
            
            # Redirect to the Google Docs URL
            google_docs_url = f'https://docs.google.com/document/d/{file.get("id")}/edit'
            return redirect(google_docs_url)
            
    except Exception as e:
        raise Exception(f"Error generating document: {str(e)}")

def process_form_data(form):
    """Processes form data and converts it into the required context format."""
    try:
        context = {
            "document_version": [],
            "category_table": [],
            "credits_table": [],
            "total_credits": clean_int(form.get("total_credits", 0)),
            "dept_Code": clean_text(form.get("dept_Code", "")),
            **{key: [] for key in ["BSC", "ESC", "PCC", "ELECTIVE", "OEC", "MC", "EEC", "HSMC"]},
            **{f"course_table{i}": [] for i in range(1, 9)}
        }

        # Process Document Version
        i = 1
        while f"document_version_version_{i}" in form:
            context["document_version"].append({
                "version": clean_text(form.get(f"document_version_version_{i}", "")),
                "date": clean_text(form.get(f"document_version_date_{i}", "")),
                "author": clean_text(form.get(f"document_version_author_{i}", "")),
                "updates": clean_text(form.get(f"document_version_updates_{i}", "")),
                "approved_by": clean_text(form.get(f"document_version_approved_{i}", ""))
            })
            i += 1
        
        context["structure_total_credits"] = sum(item["credits"] for item in context["category_table"]) 
        context["regulation"] = clean_text(form.get("reg", ""))
        context["dept"] = clean_text(form.get(f"dept", ""))
        # Process Category Table (Structure of Program)
        i = 1
        while f"structure_of_program_category_{i}" in form:
            context["category_table"].append({
                "s_no": clean_int(form.get(f"structure_of_program_sno_{i}", i)),
                "category": clean_text(form.get(f"structure_of_program_category_{i}", "")),
                "credits": clean_int(form.get(f"structure_of_program_credits_{i}", "0")),
            })
            i += 1

        # Process Credits Table (Definition of Credit)
        i = 1
        while f"definition_of_credits_l_{i}" in form:
            context["credits_table"].append({
                "l": clean_text(form.get(f"definition_of_credits_l_{i}", "")),
                "t": clean_text(form.get(f"definition_of_credits_t_{i}", "")),
                "p": clean_text(form.get(f"definition_of_credits_p_{i}", "")),
            })
            i += 1

        j = 1
        for key in ["BSC", "ESC", "PCC", "ELECTIVE", "OEC", "MC", "EEC", "HSMC"]:
            
            context[f"{key}_total_credits"] = clean_text(form.get(f"{key}_total_credits", ""))

        # Process Courses (BSC, ESC, PCC, ELECTIVE, etc.)
        for key in ["BSC", "ESC", "PCC", "ELECTIVE", "OEC", "MC", "EEC", "HSMC"]:
            i = 1
            while f"{key}_title_{i}" in form:
                context[key].append({
                    "s_no": clean_int(form.get(f"{key}_sno_{i}", i)),
                    "title": clean_text(form.get(f"{key}_title_{i}", "")),
                    "sem": clean_text(form.get(f"{key}_semester_{i}", "")),
                    "ltpc": clean_text(form.get(f"{key}_ltpc_{i}", ""))
                })
                i += 1


        for table_key in [f"course_table{i}" for i in range(1, 9)]:           
            context[f"{table_key}_total_credits"] = clean_text(form.get(f"{table_key}_total_credits", ""))
           

        # Process Course Tables (course_table1 to course_table8)
        for table_key in [f"course_table{i}" for i in range(1, 9)]:
            i = 1
            while f"{table_key}_course_code_{i}" in form:
                context[table_key].append({
                    "s_no": clean_int(form.get(f"{table_key}_sno_{i}", i)),
                    "type": clean_text(form.get(f"{table_key}_type_{i}", "")),
                    "course_code": clean_text(form.get(f"{table_key}_course_code_{i}", "")),
                    "course_title": clean_text(form.get(f"{table_key}_course_title_{i}", "")),
                    "ltpc": clean_text(form.get(f"{table_key}_ltpc_{i}", "")),
                })
                i += 1

        return context
    except Exception as e:
        logging.error(f"Error in processing form data: {e}")
        raise e  # Raise the error for debugging

def generate_docx(context):
    """Generates the Word document using docxtpl and returns it as an in-memory file."""
    template_path = "curriculum.docx"
    
    if not os.path.exists(template_path):
        logging.error("Template file 'curriculum.docx' not found.")
        raise FileNotFoundError("Template file 'curriculum.docx' not found. Please upload the correct template.")
    
    try:
        doc = DocxTemplate(template_path)
        doc.render(context)
    except Exception as e:
        logging.error(f"Template rendering error: {e}")
        raise e  # Re-raise the error for debugging

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return doc_io

@app.route('/merge_documents', methods=['POST'])
@login_required
def merge_documents():
    if current_user.role != 'teacher':
        return jsonify({'success': False, 'message': 'Access denied'})
    
    try:
        # Get subject folder
        subject_folder = DriveDirectory.query.filter_by(
            department=current_user.department,
            type='subject',
            semester=current_user.semester
        ).first()
        
        if not subject_folder:
            return jsonify({'success': False, 'message': 'Subject folder not found'})
        
        # Get Google Drive service
        drive_service, _ = get_google_drive_service()
        if not drive_service:
            return jsonify({'success': False, 'message': 'Could not connect to Google Drive'})
        
        # List all documents in subject folder
        files = drive_service.files().list(
            q=f"'{subject_folder.drive_id}' in parents and mimeType='application/vnd.google-apps.document'",
            fields="files(id, name)"
        ).execute().get('files', [])
        
        if not files:
            return jsonify({'success': False, 'message': 'No documents found to merge'})

        # Create a new merged document
        merged_doc = Document()
        
        # Download and merge each document
        for file in files:
            try:
                # Export Google Doc as docx
                doc_content = drive_service.files().export(
                    fileId=file['id'],
                    mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                ).execute()
                
                # Create a temporary file to store the downloaded document
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
                    temp_doc.write(doc_content)
                    temp_doc.flush()
                    
                    # Open the document and append its content
                    doc = Document(temp_doc.name)
                    
                    # Add a page break before each document (except the first one)
                    if merged_doc.paragraphs:
                        merged_doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
                    
                    # Add document title
                    merged_doc.add_heading(file['name'], level=1)
                    
                    # Copy all elements from the document
                    for element in doc.element.body:
                        merged_doc.element.body.append(element)
                    
                    # Clean up temp file
                    os.unlink(temp_doc.name)
            except Exception as e:
                print(f"Error processing file {file['name']}: {str(e)}")
                continue

        # Save merged document to temporary file
        temp_merged = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_merged_path = temp_merged.name
        temp_merged.close()
        merged_doc.save(temp_merged_path)

        try:
            # Upload merged document to Drive
            file_metadata = {
                'name': f'Merged_Documents_{current_user.department}_Sem{current_user.semester}',
                'parents': [subject_folder.drive_id],
                'mimeType': 'application/vnd.google-apps.document'
            }
            
            media = MediaFileUpload(
                temp_merged_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                resumable=True
            )
            
            file = drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id'
            ).execute()

            # Create approval record
            approval = DocumentApproval(
                department=current_user.department,
                semester=str(current_user.semester),
                merged_file_id=file.get('id'),
                status='pending'
            )
            db.session.add(approval)
            db.session.commit()

            return jsonify({
                'success': True,
                'message': 'Documents merged successfully',
                'file_id': file.get('id')
            })

        finally:
            # Clean up temporary merged file
            if os.path.exists(temp_merged_path):
                os.unlink(temp_merged_path)

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/approve_document/<doc_id>', methods=['POST'])
@login_required
def approve_document(doc_id):
    if current_user.role != 'advisor':
        return jsonify({'success': False, 'message': 'Access denied'})
    
    try:
        doc = DocumentApproval.query.get(doc_id)
        if not doc:
            return jsonify({'success': False, 'message': 'Document not found'})
        
        if doc.department != current_user.department or doc.semester != str(current_user.semester):
            return jsonify({'success': False, 'message': 'Access denied'})
        
        doc.status = 'approved'
        db.session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/reject_document/<doc_id>', methods=['POST'])
@login_required
def reject_document(doc_id):
    if current_user.role != 'advisor':
        return jsonify({'success': False, 'message': 'Access denied'})
    
    try:
        doc = DocumentApproval.query.get(doc_id)
        if not doc:
            return jsonify({'success': False, 'message': 'Document not found'})
        
        if doc.department != current_user.department or doc.semester != str(current_user.semester):
            return jsonify({'success': False, 'message': 'Access denied'})
        
        doc.status = 'rejected'
        db.session.commit()
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/upload_document', methods=['POST'])
@login_required
def upload_document():
    logging.info('Upload document route called')
    if current_user.role != 'teacher':
        logging.warning('Access denied for user: %s', current_user.username)
        return jsonify({'success': False, 'message': 'Access denied'})
    
    if 'file' not in request.files:
        logging.warning('No file uploaded')
        return jsonify({'success': False, 'message': 'No file uploaded'})
    
    file = request.files['file']
    if file.filename == '':
        logging.warning('No file selected')
        return jsonify({'success': False, 'message': 'No file selected'})
    
    temp_dir = tempfile.mkdtemp()
    temp_file_path = None
    
    try:
        # Get Google Drive service
        drive_service, auth_url = get_google_drive_service()
        if not drive_service:
            if auth_url:
                # Store the current request in session
                session['pending_upload'] = {
                    'filename': file.filename,
                    'department': current_user.department,
                    'semester': current_user.semester
                }
                logging.info('Redirecting to auth_url for authentication')
                return jsonify({'success': False, 'auth_url': auth_url})
            else:
                logging.error('Could not connect to Google Drive and no auth URL provided')
                return jsonify({'success': False, 'message': 'Could not connect to Google Drive'})
        
        # Get subject folder
        subject_folder = DriveDirectory.query.filter_by(
            department=current_user.department,
            type='subject',
            semester=current_user.semester
        ).first()
        
        if not subject_folder:
            logging.error('Subject folder not found')
            return jsonify({'success': False, 'message': 'Subject folder not found'})
        
        # Save file to temporary directory
        temp_file_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(temp_file_path)
        logging.info('File saved temporarily at %s', temp_file_path)
        
        # Upload the file to Google Drive
        file_metadata = {
            'name': file.filename,
            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'parents': [subject_folder.drive_id]
        }
        
        media = MediaFileUpload(
            temp_file_path,
            mimetype=file.content_type,
            resumable=True
        )
        
        uploaded_file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        
        logging.info('File uploaded successfully with ID: %s', uploaded_file['id'])
        
        # Create approval record
        approval = DocumentApproval(
            department=current_user.department,
            semester=str(current_user.semester),
            merged_file_id=uploaded_file['id'],
            document_name=file.filename,
            status='pending'
        )
        db.session.add(approval)
        db.session.commit()
        logging.info('Approval record created for file ID: %s', uploaded_file['id'])
        
        return jsonify({'success': True, 'message': 'File uploaded successfully'})
    
    except Exception as e:
        logging.error('Error occurred: %s', str(e))
        return jsonify({'success': False, 'message': str(e)})
    
    finally:
        # Clean up
        try:
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            if temp_dir and os.path.exists(temp_dir):
                os.rmdir(temp_dir)
            logging.info('Temporary files cleaned up successfully')
        except Exception as e:
            logging.warning('Failed to clean up temporary files: %s', str(e))

@app.route('/upload_hod_document', methods=['POST'])
@login_required
def upload_hod_document():
    if current_user.role != 'advisor':
        return jsonify({'success': False, 'message': 'Access denied'})
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file uploaded'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No file selected'})
    
    temp_dir = tempfile.mkdtemp()
    temp_file_path = None
    
    try:
        drive_service, auth_url = get_google_drive_service()
        if not drive_service:
            if auth_url:
                session['pending_upload'] = {
                    'filename': file.filename,
                    'department': current_user.department,
                    'semester': current_user.semester,
                    'document_type': 'hod'
                }
                return jsonify({'success': False, 'auth_url': auth_url})
            else:
                return jsonify({'success': False, 'message': 'Could not connect to Google Drive'})
        
        # Get HOD folder
        hod_folder = DriveDirectory.query.filter_by(
            department=current_user.department,
            type='hod',
            semester=current_user.semester
        ).first()
        
        if not hod_folder:
            return jsonify({'success': False, 'message': 'HOD folder not found'})
        
        temp_file_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(temp_file_path)
        
        file_metadata = {
            'name': file.filename,
            'parents': [hod_folder.drive_id]
        }
        
        media = MediaFileUpload(
            temp_file_path,
            mimetype=file.content_type,
            resumable=True
        )
        
        uploaded_file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        
        approval = DocumentApproval(
            department=current_user.department,
            semester=str(current_user.semester),
            merged_file_id=uploaded_file['id'],
            document_name=file.filename,
            document_type='hod',
            status='pending'
        )
        db.session.add(approval)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'File uploaded successfully'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})
    
    finally:
        try:
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            if temp_dir and os.path.exists(temp_dir):
                os.rmdir(temp_dir)
        except Exception as e:
            logging.warning('Failed to clean up temporary files: %s', str(e))

@app.route('/upload_syllabus', methods=['POST'])
@login_required
def upload_syllabus():
    if current_user.role != 'advisor':
        return jsonify({'success': False, 'message': 'Access denied'})
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file uploaded'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No file selected'})
    
    temp_dir = tempfile.mkdtemp()
    temp_file_path = None
    
    try:
        drive_service, auth_url = get_google_drive_service()
        if not drive_service:
            if auth_url:
                session['pending_upload'] = {
                    'filename': file.filename,
                    'department': current_user.department,
                    'semester': current_user.semester,
                    'document_type': 'syllabus'
                }
                return jsonify({'success': False, 'auth_url': auth_url})
            else:
                return jsonify({'success': False, 'message': 'Could not connect to Google Drive'})
        
        # Get syllabus folder
        syllabus_folder = DriveDirectory.query.filter_by(
            department=current_user.department,
            type='syllabus',
            semester=current_user.semester
        ).first()
        
        if not syllabus_folder:
            return jsonify({'success': False, 'message': 'Syllabus folder not found'})
        
        temp_file_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(temp_file_path)
        
        file_metadata = {
            'name': file.filename,
            'parents': [syllabus_folder.drive_id]
        }
        
        media = MediaFileUpload(
            temp_file_path,
            mimetype=file.content_type,
            resumable=True
        )
        
        uploaded_file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        
        approval = DocumentApproval(
            department=current_user.department,
            semester=str(current_user.semester),
            merged_file_id=uploaded_file['id'],
            document_name=file.filename,
            document_type='syllabus',
            status='pending'
        )
        db.session.add(approval)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'File uploaded successfully'})
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})
    
    finally:
        try:
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            if temp_dir and os.path.exists(temp_dir):
                os.rmdir(temp_dir)
        except Exception as e:
            logging.warning('Failed to clean up temporary files: %s', str(e))

@app.route('/approve_syllabus/<int:doc_id>', methods=['POST'])
@login_required
def approve_syllabus(doc_id):
    if current_user.role != 'advisor':
        return jsonify({'success': False, 'message': 'Access denied'})
    
    doc = DocumentApproval.query.get_or_404(doc_id)
    if doc.department != current_user.department or doc.semester != str(current_user.semester):
        return jsonify({'success': False, 'message': 'Access denied'})
    
    doc.status = 'approved'
    doc.approved_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Document approved successfully'})

@app.route('/reject_syllabus/<int:doc_id>', methods=['POST'])
@login_required
def reject_syllabus(doc_id):
    if current_user.role != 'advisor':
        return jsonify({'success': False, 'message': 'Access denied'})
    
    doc = DocumentApproval.query.get_or_404(doc_id)
    if doc.department != current_user.department or doc.semester != str(current_user.semester):
        return jsonify({'success': False, 'message': 'Access denied'})
    
    doc.status = 'rejected'
    doc.rejected_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Document rejected successfully'})

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
